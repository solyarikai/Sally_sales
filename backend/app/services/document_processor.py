"""
Document Processor - Convert various file formats to Markdown for AI
"""
import io
from typing import List, Optional
from openai import AsyncOpenAI
from app.core.config import settings

client = AsyncOpenAI(api_key=settings.OPENAI_API_KEY)


async def process_document(content: bytes, filename: str, content_type: Optional[str] = None) -> dict:
    """
    Process uploaded document and convert to markdown.
    Supports: PDF, DOCX, TXT, and other text formats.
    """
    raw_text = ""
    
    # Determine file type
    filename_lower = filename.lower()
    
    if filename_lower.endswith('.pdf'):
        raw_text = await extract_pdf_text(content)
    elif filename_lower.endswith('.docx'):
        raw_text = await extract_docx_text(content)
    elif filename_lower.endswith(('.txt', '.md', '.csv', '.json')):
        raw_text = content.decode('utf-8', errors='ignore')
    elif content_type and 'text' in content_type:
        raw_text = content.decode('utf-8', errors='ignore')
    else:
        # Try to decode as text
        try:
            raw_text = content.decode('utf-8', errors='ignore')
        except:
            raise ValueError(f"Unsupported file type: {filename}")
    
    if not raw_text or len(raw_text.strip()) < 10:
        raise ValueError("Could not extract text from document")
    
    # Skip AI conversion for files already in markdown format
    if filename_lower.endswith('.md'):
        return {
            "raw_text": raw_text,
            "markdown": raw_text  # Already markdown, use as-is
        }
    
    # Convert to markdown using AI
    markdown = await convert_to_markdown(raw_text, filename)
    
    return {
        "raw_text": raw_text,
        "markdown": markdown
    }


async def extract_pdf_text(content: bytes) -> str:
    """Extract text from PDF using pypdf"""
    try:
        from pypdf import PdfReader
        
        pdf_file = io.BytesIO(content)
        reader = PdfReader(pdf_file)
        
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        return "\n\n".join(text_parts)
    except ImportError:
        # Fallback - try pypdf2
        try:
            from PyPDF2 import PdfReader
            
            pdf_file = io.BytesIO(content)
            reader = PdfReader(pdf_file)
            
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            return "\n\n".join(text_parts)
        except ImportError:
            raise ValueError("PDF support not available. Install pypdf: pip install pypdf")


async def extract_docx_text(content: bytes) -> str:
    """Extract text from DOCX"""
    try:
        from docx import Document
        
        docx_file = io.BytesIO(content)
        doc = Document(docx_file)
        
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return "\n\n".join(text_parts)
    except ImportError:
        raise ValueError("DOCX support not available. Install python-docx: pip install python-docx")


async def convert_to_markdown(text: str, filename: str) -> str:
    """Convert raw text to clean markdown using AI"""
    # Truncate if too long
    max_chars = 50000
    if len(text) > max_chars:
        text = text[:max_chars] + "\n\n[... truncated ...]"
    
    try:
        response = await client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": """Convert this document to clean, well-structured Markdown.

Guidelines:
- Keep ALL important information
- Use appropriate headers (# ## ###) for sections
- Use bullet points and numbered lists where appropriate
- Format tables properly if data is tabular
- Remove noise like page numbers, headers/footers
- Make it easy for an AI to understand and reference
- Preserve key facts, numbers, names, and details

Return ONLY the markdown content, no explanations or meta-comments."""
                },
                {
                    "role": "user",
                    "content": f"Filename: {filename}\n\n---\n\n{text}"
                }
            ],
            temperature=0.1,
            max_tokens=4000
        )
        
        return response.choices[0].message.content.strip()
        
    except Exception as e:
        # Fallback - return raw text with basic formatting
        return f"# {filename}\n\n{text}"


async def generate_company_summary(documents_md: List[str]) -> str:
    """Generate company summary from all documents"""
    if not documents_md:
        return ""
    
    # Combine all documents
    combined = "\n\n---\n\n".join(documents_md)
    
    # Truncate if too long
    max_chars = 80000
    if len(combined) > max_chars:
        combined = combined[:max_chars] + "\n\n[... additional content truncated ...]"
    
    try:
        response = await client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": """Create a comprehensive company summary from these documents.

The summary should include:
1. **Company Overview** - What the company does, core business
2. **Products/Services** - What they offer
3. **Target Market** - Who they serve
4. **Key Differentiators** - What makes them unique
5. **Social Proof** - Notable clients, metrics, achievements
6. **Contact/Resources** - Key links, booking info if mentioned

Format as clean Markdown. Be comprehensive but concise.
This summary will be used as context for AI-powered outreach, so include details useful for personalization."""
                },
                {
                    "role": "user",
                    "content": f"Documents:\n\n{combined}"
                }
            ],
            temperature=0.2,
            max_tokens=2000
        )
        
        return response.choices[0].message.content.strip()
        
    except Exception as e:
        return f"Error generating summary: {str(e)}"
